#!/usr/bin/env python
"""Lightweight Markdown -> styled .docx (python-docx).
Handles: # / ## / ### headings, --- rules, - bullets, > blockquotes,
pipe tables, **bold**, `code`, and paragraphs. Brand-styled (cream/gold).
Usage: python scripts/md-to-docx.py <in.md> <out.docx>
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

GOLD = RGBColor(0xB7, 0x92, 0x4E)
DARK = RGBColor(0x2C, 0x24, 0x16)

def add_runs(par, text):
    # split on **bold** and `code`, keep delimiters
    for chunk in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            r = par.add_run(chunk[2:-2]); r.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = par.add_run(chunk[1:-1]); r.font.name = "Consolas"
        else:
            par.add_run(chunk)

def build(md_path, docx_path):
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(11); n.font.color.rgb = DARK
    lines = open(md_path, encoding="utf-8").read().splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()
        if not ln.strip():
            i += 1; continue
        if ln.startswith("# "):
            p = doc.add_paragraph(); r = p.add_run(ln[2:].replace("**","")); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=DARK
        elif ln.startswith("## "):
            p = doc.add_paragraph(); r = p.add_run(ln[3:].replace("**","")); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=GOLD
        elif ln.startswith("### "):
            p = doc.add_paragraph(); r = p.add_run(ln[4:].replace("**","")); r.bold=True; r.font.size=Pt(12.5); r.font.color.rgb=DARK
        elif ln.strip() == "---":
            doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif ln.startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, ln.lstrip("> ").strip());
            for r in p.runs: r.italic = True
        elif ln.lstrip().startswith("- "):
            while i < len(lines) and lines[i].lstrip().startswith("- "):
                p = doc.add_paragraph(style="List Bullet"); add_runs(p, lines[i].lstrip()[2:]); i += 1
            continue
        elif ln.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[\s:\-]+$", "".join(cells)):
                    rows.append(cells)
                i += 1
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Light Grid Accent 4"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        cp = t.cell(ri, ci).paragraphs[0]; add_runs(cp, cell)
                        if ri == 0:
                            for r in cp.runs: r.bold = True
            continue
        else:
            p = doc.add_paragraph(); add_runs(p, ln)
        i += 1
    doc.save(docx_path)
    print("wrote", docx_path)

if __name__ == "__main__":
    build(sys.argv[1], sys.argv[2])
